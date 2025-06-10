"""Aplicação para geração de contratos e envio por e-mail."""

import os
from email.message import EmailMessage
from io import BytesIO
import smtplib
from flask import Flask, render_template, request
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "Contrato Vitorino.docx")

FORM_FIELDS = {
    "Comprador": "Comprador",
    "CPF": "CPF",
    "RG": "RG",
    "Emissor": "Emissor",
    "EstadoCivil": "EstadoCivil",
    "Profissao": "Profissão",
    "Endereco": "Endereço",
    "Numero": "Número",
    "Complemento": "Complemento",
    "Bairro": "Bairro",
    "Cidade": "Cidade",
    "CEP": "CEP",
    "Quadra": "Quadra",
    "Lote": "Lote",
    "Testemunha": "Testemunha",
    "CPFTest": "CPF Test",
    "Testemunha2": "Testemunha2",
    "CPFTest2": "CPF Test2",
}

app = Flask(__name__)

def _replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a single paragraph preserving formatting."""
    if not paragraph.runs:
        return

    # Join all text to detect placeholders across multiple runs
    full_text = ''.join(run.text for run in paragraph.runs)
    replaced = False
    for key, val in replacements.items():
        placeholder = f'[{key}]'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, val)
            replaced = True

    if replaced:
        # Write merged text back keeping the formatting of the first run
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        # Fallback: simple in-run replacement
        for run in paragraph.runs:
            for key, val in replacements.items():
                placeholder = f'[{key}]'
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, val)

def _process_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested in cell.tables:
                _process_table(nested, replacements)


def replace_placeholders(replacements):
    """Gera um novo DOCX substituindo marcadores do template de forma segura."""
    doc = Document(TEMPLATE_PATH)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _process_table(table, replacements)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def send_email(doc_bytes):
    user = os.environ.get('EMAIL_USER')
    password = os.environ.get('EMAIL_PASS')
    dest = os.environ.get('EMAIL_DEST', 'rba1807@gmail.com')
    if not user or not password:
        raise RuntimeError('Credenciais de e-mail não definidas')
    msg = EmailMessage()
    msg['Subject'] = 'Contrato Gerado'
    msg['From'] = user
    msg['To'] = dest
    msg.set_content('Segue contrato em anexo.')
    msg.add_attachment(
        doc_bytes,
        maintype='application',
        subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename='contrato.docx'
    )
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(user, password)
        smtp.send_message(msg)


@app.route('/', methods=['GET', 'POST'])
def form():
    status = None
    if request.method == 'POST':
        data = {f: request.form.get(f, '') for f in FORM_FIELDS.keys()}
        replacements = {placeholder: data[field] for field, placeholder in FORM_FIELDS.items()}
        try:
            doc = replace_placeholders(replacements)
            send_email(doc)
            status = 'Contrato enviado com sucesso!'
        except Exception as exc:  # pragma: no cover - best effort
            print('Erro:', exc)
            status = 'Falha ao enviar contrato.'
    return render_template('form.html', status=status)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8000"))
    app.run(host="0.0.0.0", port=port)
